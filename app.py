import streamlit as st
import os
from groq import Groq
import google.generativeai as genai
from docx import Document
import io
from pydub import AudioSegment
import math

# Nastavení vzhledu
st.set_page_config(page_title="Chytrý zápis ze schůzky", page_icon="📝", layout="centered")

# Načtení klíčů
try:
    groq_api_key = st.secrets["GROQ_API_KEY"]
    gemini_api_key = st.secrets["GEMINI_API_KEY"]
except:
    st.error("Chybí API klíče v nastavení aplikace.")
    st.stop()

# --- PAMĚŤ APLIKACE ---
if "transcription" not in st.session_state:
    st.session_state.transcription = None
if "zapis_text" not in st.session_state:
    st.session_state.zapis_text = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

st.title("📝 Generátor manažerských zápisů")
st.markdown("Nahraj audio ze schůzky. Aplikace si sama poradí i s velkými soubory (automaticky si je rozseká).")

audio_file = st.file_uploader("Nahraj záznam ze schůzky (MP3, WAV, M4A)", type=['mp3', 'wav', 'm4a'])

col_btn1, col_btn2 = st.columns(2)
with col_btn1:
    btn_standard = st.button("🚀 Vygenerovat standardní zápis", use_container_width=True)
with col_btn2:
    btn_obecny = st.button("📋 Vygenerovat obecný zápis", use_container_width=True)

if btn_standard or btn_obecny:
    if not audio_file:
        st.warning("Nejprve prosím nahraj soubor s audiem.")
    else:
        try:
            st.session_state.chat_history = []
            
            # --- CHYTRÉ ZPRACOVÁNÍ AUDIA ---
            file_extension = audio_file.name.split('.')[-1].lower()
            temp_filename = f"temp_original.{file_extension}"
            
            # Uložení originálu
            with open(temp_filename, "wb") as f:
                f.write(audio_file.getbuffer())
            
            # Zjištění velikosti
            file_size_mb = os.path.getsize(temp_filename) / (1024 * 1024)
            client = Groq(api_key=groq_api_key)
            full_transcription = ""

            # Pokud je soubor větší než 24 MB, rozsekáme ho na 10minutové úseky
            if file_size_mb > 24:
                st.info(f"Soubor je velký ({file_size_mb:.1f} MB). Rozděluji ho na menší části. To může chvilku trvat...")
                audio = AudioSegment.from_file(temp_filename)
                
                # 10 minut = 600 000 milisekund
                chunk_length_ms = 10 * 60 * 1000 
                chunks_count = math.ceil(len(audio) / chunk_length_ms)
                
                for i in range(chunks_count):
                    start_time = i * chunk_length_ms
                    end_time = (i + 1) * chunk_length_ms
                    chunk = audio[start_time:end_time]
                    
                    chunk_filename = f"chunk_{i}.mp3"
                    chunk.export(chunk_filename, format="mp3")
                    
                    with st.spinner(f"⏳ Poslouchám část {i+1} z {chunks_count}..."):
                        with open(chunk_filename, "rb") as file:
                            transcription = client.audio.transcriptions.create(
                                file=(chunk_filename, file.read()),
                                model="whisper-large-v3",
                                response_format="text",
                                language="cs"
                            )
                            full_transcription += transcription + " "
                    
                    os.remove(chunk_filename) # Úklid kousku
            else:
                # Běžný malý soubor
                with st.spinner("⏳ Poslouchám a přepisuji audio..."):
                    with open(temp_filename, "rb") as file:
                        transcription = client.audio.transcriptions.create(
                          file=(temp_filename, file.read()),
                          model="whisper-large-v3",
                          response_format="text",
                          language="cs"
                        )
                        full_transcription = transcription

            os.remove(temp_filename) # Úklid originálu
            st.session_state.transcription = full_transcription
            st.success("✅ Přepis byl úspěšně dokončen!")

            # --- TVORBA ZÁPISU ---
            with st.spinner("⏳ Generuji zápis podle vybrané šablony..."):
                genai.configure(api_key=gemini_api_key)
                model = genai.GenerativeModel('gemini-2.5-flash')
                
                if btn_standard:
                    prompt = f"""
                    Jsi profesionální firemní asistent. Přečti si následující surový přepis ze schůzky a vytvoř z něj přehledný manažerský zápis v češtině.
                    Rozděl ho na:
                    1. Hlavní téma schůzky
                    2. Nejdůležitější probrané body (odrážky)
                    3. Učiněná rozhodnutí
                    4. Akční kroky / Úkoly (Kdo má co udělat)
                    
                    Zde je přepis:
                    {st.session_state.transcription}
                    """
                elif btn_obecny:
                    prompt = f"""
                    Jsi profesionální firemní asistent. Přečti si následující surový přepis ze schůzky a vytvoř z něj přesný zápis v češtině PŘESNĚ podle následující šablony. 
                    Dodržuj formátování (nadpisy, tučné písmo) a řiď se instrukcemi, které jsou uvedeny v hranatých závorkách.

                    ## MANAZERSKE SHRNUTÍ
                    **Cíl setkání:** [jedna az dve vety]
                    **Klícová rozhodnutí:** [kazde rozhodnutí na novy radek s pomlckou;pokud zadne nepadlo napís: Bez formalnich rozhodnutí]
                    ---
                    ## DISKUTOVANÁ TÉMATA
                    [kazde tema na novy radek s pomlckou, max 8 bodu]
                    ---
                    ## AKCNÍ KROKY
                    | # | Úkol | Odpovědná osoba | Termín | Stav |
                    |---|------|-----------------|--------|------|
                    [radky tabulky; pokud neni termin nebo osoba napís Neurčeno; Stav vzdy Nový]

                    Zde je přepis:
                    {st.session_state.transcription}
                    """

                response = model.generate_content(prompt)
                st.session_state.zapis_text = response.text
                
        except Exception as e:
            st.error(f"Ouvej, něco se pokazilo: {e}")

# --- ZOBRAZENÍ VÝSLEDKŮ A CHATU ---
if st.session_state.transcription and st.session_state.zapis_text:
    
    st.success("✅ Zápis je hotový!")
    st.markdown("### Náhled zápisu:")
    st.write(st.session_state.zapis_text)

    st.markdown("### 💾 Ke stažení:")
    col1, col2 = st.columns(2)
    
    with col1:
        doc_zapis = Document()
        doc_zapis.add_heading('Zápis ze schůzky', 0)
        doc_zapis.add_paragraph(st.session_state.zapis_text)
        bio_zapis = io.BytesIO()
        doc_zapis.save(bio_zapis)
        
        st.download_button(
            label="📝 Stáhnout zápis",
            data=bio_zapis.getvalue(),
            file_name="zapis_ze_schuzky.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    with col2:
        doc_prepis = Document()
        doc_prepis.add_heading('Kompletní přepis schůzky', 0)
        doc_prepis.add_paragraph(st.session_state.transcription)
        bio_prepis = io.BytesIO()
        doc_prepis.save(bio_prepis)
        
        st.download_button(
            label="🗣️ Stáhnout doslovný přepis",
            data=bio_prepis.getvalue(),
            file_name="kompletni_prepis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.markdown("---")
    st.markdown("### 💬 Zeptejte se na detaily ze schůzky")
    
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if user_question := st.chat_input("Zeptejte se... (např. 'Jaký byl dohodnutý termín spuštění?')"):
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        with st.chat_message("user"):
            st.markdown(user_question)

        with st.chat_message("assistant"):
            with st.spinner("Dohledávám v přepisu..."):
                genai.configure(api_key=gemini_api_key)
                model = genai.GenerativeModel('gemini-2.5-flash')
                
                chat_prompt = f"""
                Jsi asistent. Tvojí jedinou prací je odpovídat na otázky týkající se této schůzky, POUZE na základě poskytnutého přepisu.
                Pokud odpověď v přepisu nenajdeš, omluv se a řekni: "Tato informace v přepisu bohužel nezazněla."
                
                Zde je přepis schůzky:
                {st.session_state.transcription}
                
                Zde je otázka uživatele:
                {user_question}
                """
                
                odpoved = model.generate_content(chat_prompt)
                st.markdown(odpoved.text)
                st.session_state.chat_history.append({"role": "assistant", "content": odpoved.text})
